"""
Document loaders for extracting text from various file formats.
Uses LangChain document loaders with custom extensions for OCR.
"""

import os
from typing import List, Dict, Any, Optional
from pathlib import Path

from langchain_core.documents import Document

# Try to import LangChain community loaders
try:
    from langchain_community.document_loaders import (
        PyMuPDFLoader,
        Docx2txtLoader,
        TextLoader,
    )
    LANGCHAIN_LOADERS_AVAILABLE = True
except ImportError:
    LANGCHAIN_LOADERS_AVAILABLE = False
    print("Warning: langchain-community not installed. Using fallback loaders.")


class ImageLoader:
    """
    Custom image loader using OCR.
    Extracts text from images using Tesseract OCR.
    """
    
    def __init__(self, file_path: str):
        """
        Initialize image loader.
        
        Args:
            file_path: Path to the image file
        """
        self.file_path = file_path
    
    def load(self) -> List[Document]:
        """
        Load and extract text from image using OCR.
        
        Returns:
            List containing a single Document with OCR text
        """
        from ingestion.ocr import OCRProcessor
        
        if not os.path.exists(self.file_path):
            raise FileNotFoundError(f"Image file not found: {self.file_path}")
        
        ocr_processor = OCRProcessor()
        result = ocr_processor.extract_with_metadata(self.file_path)
        
        return [Document(
            page_content=result['text'],
            metadata={
                'source': str(self.file_path),
                'filename': result['filename'],
                'image_size': result.get('image_size'),
                'type': 'image'
            }
        )]


class PDFLoaderWithOCR:
    """
    PDF loader with OCR fallback for scanned/image-based PDFs.
    Uses PyMuPDF for text extraction and falls back to OCR when needed.
    """
    
    def __init__(self, file_path: str, ocr_fallback: bool = True):
        """
        Initialize PDF loader.
        
        Args:
            file_path: Path to the PDF file
            ocr_fallback: Whether to use OCR for pages with no text
        """
        self.file_path = file_path
        self.ocr_fallback = ocr_fallback
    
    def load(self) -> List[Document]:
        """
        Load and extract text from PDF.
        
        Returns:
            List of Documents, one per page
        """
        import fitz  # PyMuPDF
        
        if not os.path.exists(self.file_path):
            raise FileNotFoundError(f"PDF file not found: {self.file_path}")
        
        documents = []
        filename = Path(self.file_path).name
        
        try:
            doc = fitz.open(self.file_path)
            total_pages = len(doc)
            
            for page_num in range(total_pages):
                page = doc[page_num]
                page_text = page.get_text("text")
                
                # If no text and OCR fallback enabled, try OCR
                if not page_text.strip() and self.ocr_fallback:
                    page_text = self._ocr_page(page)
                
                if page_text.strip():
                    documents.append(Document(
                        page_content=page_text,
                        metadata={
                            'source': str(self.file_path),
                            'filename': filename,
                            'page': page_num + 1,
                            'total_pages': total_pages,
                            'type': 'pdf'
                        }
                    ))
            
            doc.close()
            
        except Exception as e:
            raise RuntimeError(f"Error reading PDF file: {e}")
        
        return documents
    
    def _ocr_page(self, page) -> str:
        """OCR a PDF page by rendering it as an image."""
        try:
            import fitz
            from ingestion.ocr import OCRProcessor
            
            # Render page to image
            mat = fitz.Matrix(2, 2)  # 2x zoom
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")
            
            # OCR the rendered page
            ocr_processor = OCRProcessor()
            return ocr_processor.extract_text_from_bytes(img_bytes, preprocess=True)
        except Exception:
            return ""


class UnifiedDocumentLoader:
    """
    Unified document loader that automatically selects the appropriate loader
    based on file extension. Uses LangChain loaders where available.
    """
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.png', '.jpg', '.jpeg'}
    IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg'}
    
    def __init__(self):
        """Initialize unified loader."""
        pass
    
    def _get_loader(self, file_path: str):
        """Get the appropriate loader for a file."""
        extension = Path(file_path).suffix.lower()
        
        if extension == '.pdf':
            if LANGCHAIN_LOADERS_AVAILABLE:
                # Use custom PDF loader with OCR fallback
                return PDFLoaderWithOCR(file_path)
            else:
                return PDFLoaderWithOCR(file_path)
        
        elif extension == '.docx':
            if LANGCHAIN_LOADERS_AVAILABLE:
                return Docx2txtLoader(file_path)
            else:
                return self._fallback_docx_loader(file_path)
        
        elif extension == '.txt':
            if LANGCHAIN_LOADERS_AVAILABLE:
                return TextLoader(file_path, encoding='utf-8')
            else:
                return self._fallback_text_loader(file_path)
        
        elif extension in self.IMAGE_EXTENSIONS:
            return ImageLoader(file_path)
        
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    def _fallback_docx_loader(self, file_path: str):
        """Fallback DOCX loader."""
        class _DocxLoader:
            def __init__(self, path):
                self.path = path
            
            def load(self):
                from docx import Document as DocxDocument
                doc = DocxDocument(self.path)
                text_parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            text_parts.append(row_text)
                return [Document(
                    page_content="\n\n".join(text_parts),
                    metadata={'source': self.path, 'filename': Path(self.path).name, 'type': 'docx'}
                )]
        return _DocxLoader(file_path)
    
    def _fallback_text_loader(self, file_path: str):
        """Fallback text loader."""
        class _TextLoader:
            def __init__(self, path):
                self.path = path
            
            def load(self):
                encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
                for encoding in encodings:
                    try:
                        with open(self.path, 'r', encoding=encoding) as f:
                            content = f.read()
                        return [Document(
                            page_content=content,
                            metadata={'source': self.path, 'filename': Path(self.path).name, 'type': 'txt'}
                        )]
                    except UnicodeDecodeError:
                        continue
                raise RuntimeError("Could not decode text file")
        return _TextLoader(file_path)
    
    def load(self, file_path: str) -> List[Document]:
        """
        Load a document.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of Document objects
        """
        loader = self._get_loader(file_path)
        return loader.load()
    
    def load_as_dict(self, file_path: str) -> Dict[str, Any]:
        """
        Load a document and return as dictionary (for backward compatibility).
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing text and metadata
        """
        docs = self.load(file_path)
        
        # Combine all documents into one text
        combined_text = "\n\n".join(doc.page_content for doc in docs)
        
        path = Path(file_path)
        return {
            'text': combined_text,
            'source': str(file_path),
            'filename': path.name,
            'extension': path.suffix.lower(),
            'documents': docs  # Include original documents
        }
    
    def load_multiple(self, file_paths: List[str]) -> List[Document]:
        """
        Load multiple documents.
        
        Args:
            file_paths: List of file paths to load
            
        Returns:
            List of all Document objects
        """
        all_docs = []
        for file_path in file_paths:
            try:
                docs = self.load(file_path)
                all_docs.extend(docs)
            except Exception as e:
                print(f"Warning: Failed to load {file_path}: {e}")
        
        return all_docs
    
    def supported_extensions(self) -> List[str]:
        """Return all supported file extensions."""
        return list(self.SUPPORTED_EXTENSIONS)


# Backward compatibility aliases
class DocumentLoader(UnifiedDocumentLoader):
    """Alias for backward compatibility."""
    
    def load(self, file_path: str) -> Dict[str, Any]:
        """Load document and return as dict (legacy API)."""
        return self.load_as_dict(file_path)


# Legacy loader classes for backward compatibility
class BaseLoader:
    """Abstract base class for document loaders (legacy)."""
    pass


class PDFLoader(PDFLoaderWithOCR):
    """Alias for backward compatibility."""
    pass


class DOCXLoader:
    """Legacy DOCX loader."""
    
    def load(self, file_path: str) -> str:
        loader = UnifiedDocumentLoader()
        result = loader.load_as_dict(file_path)
        return result['text']
    
    def supported_extensions(self) -> List[str]:
        return ['.docx']


class TXTLoader:
    """Legacy TXT loader."""
    
    def load(self, file_path: str) -> str:
        loader = UnifiedDocumentLoader()
        result = loader.load_as_dict(file_path)
        return result['text']
    
    def supported_extensions(self) -> List[str]:
        return ['.txt']
